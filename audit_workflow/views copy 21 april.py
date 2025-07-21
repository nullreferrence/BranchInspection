from django.contrib.auth import authenticate, login
from django.core.paginator import Paginator
from django.forms import modelformset_factory, inlineformset_factory, formset_factory
from django.utils.timezone import now
from .forms import AuditSubmissionForm, CommentForm, Personel_DecisionForm, PersonForm, ItemForm
from .models import AuditSubmission, Items, AuditObjection, AuditItemStatus, Branch, User, Person, Comment, Upload, \
    Personel_Decision
from .forms import AuditObjectionForm, LoginForm, UploadForm
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
from django.contrib import messages
from django.db import IntegrityError
from django.contrib.auth import logout, update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm
import re
from django.utils.html import strip_tags  # Removes HTML tags
from html.parser import HTMLParser  # Converts entities like &nbsp; to space
from docx import Document


@login_required
def start_audit(request):
    if request.method == "POST":
        form = AuditSubmissionForm(request.POST)
        all_branches = Branch.objects.all().distinct()  # Correct: distinct() added
        if form.is_valid():
            audit = form.save(commit=False)
            audit.auditor = request.user
            audit.save()
            for item in Items.objects.all():
                AuditItemStatus.objects.create(audit=audit, item=item)
            return redirect('audit_workflow:audit_summary', audit_id=audit.id)
    else:
        form = AuditSubmissionForm()
        all_branches = Branch.objects.all().distinct()  # Correct: distinct() added for GET too

    return render(request, 'audit_workflow/start_audit.html', {'form': form, 'all_branches': all_branches})  # Correct variable name


@login_required()
def audit_summary(request, audit_id):
    try:
        audit = AuditSubmission.objects.get(id=audit_id, auditor=request.user)
    except AuditSubmission.DoesNotExist:
        return render(request, 'audit_workflow/404.html', status=404)

    items = Items.objects.filter(audit_type=audit.audit_type).order_by('id')
    paginator = Paginator(items, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    item_statuses = {status.item.id: status.is_completed for status in AuditItemStatus.objects.filter(audit=audit)}
    selected_item = None
    form_html = ""

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    convicted_persons_by_item = {
        item.id: Person.objects.filter(audit_id=audit, items=item)
        for item in items
    }

    if request.method == "GET" and 'item_id' in request.GET:
        selected_item_id = request.GET.get('item_id')
        selected_item = get_object_or_404(Items, id=selected_item_id)
        objection_form = AuditObjectionForm(initial={'items': selected_item})
        upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set')  # Set prefix here
        decision_form = Personel_DecisionForm()
        person_formset = PersonFormSet(queryset=Person.objects.filter(audit_id=audit, items=selected_item), prefix='person_set') # Set prefix here

        form_html = render_to_string(
            "audit_workflow/submit_objection_form.html",
            {
                "objection_form": objection_form,
                "upload_formset": upload_formset,  # Pass UploadFormSet to the template
                "audit": audit,
                "item": selected_item,
                "decision_form": decision_form,
                "person_formset": person_formset,
            },
            request=request,
        )
        return HttpResponse(form_html)

    # Initialize forms for the main context
    decision_form = Personel_DecisionForm()
    person_formset = PersonFormSet(queryset=Person.objects.none(), prefix='person_set') # Set prefix here
    upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set') # Set prefix here
    # Items in the current page
    current_page_items = page_obj.object_list

    # Get objections only for items in the current page
    completed_objections = AuditObjection.objects.filter(submission=audit, items__in=current_page_items).select_related(
        'items')

    context = {
        "audit": audit,
        "items": page_obj,
        "item_statuses": item_statuses,
        "selected_item": selected_item,
        "form_html": form_html,
        "decision_form": decision_form,
        "person_formset": person_formset,
        'page_obj': page_obj,
        "upload_formset": upload_formset, # Pass UploadFormSet to the main context
        "convicted_persons_by_item": convicted_persons_by_item,
        "completed_objections": completed_objections,
    }
    return render(request, "audit_workflow/audit_summary.html", context)












@login_required
def finalize_audit(request, audit_id):
    audit = get_object_or_404(AuditSubmission, id=audit_id, auditor=request.user)
    incomplete_items = AuditItemStatus.objects.filter(audit=audit, is_completed=False)

    if incomplete_items.exists():
        return render(request, "audit_workflow/finalize_error.html", {"audit": audit})

    audit.submitted = True
    audit.submission_date = now()
    audit.save()
    return redirect("audit_workflow:dashboard")



@login_required
def dashboard(request):
    user = request.user
    completed_audits = []
    incomplete_audits = []

    if user.role == 'admin':
        all_audits = AuditSubmission.objects.all().order_by('-submission_date')

        for audit in all_audits:
            all_items_completed = all(status.is_completed for status in AuditItemStatus.objects.filter(audit=audit))
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)

    elif user.role == 'manager':
        # Audits related to users in the same region
        region_users = User.objects.filter(region=user.region)
        all_region_audits = AuditSubmission.objects.filter(auditor__in=region_users).order_by('-submission_date')

        for audit in all_region_audits:
            all_items_completed = all(status.is_completed for status in AuditItemStatus.objects.filter(audit=audit))
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)

    else:  # Regular user
        completed_audits = AuditSubmission.objects.filter(
            branch=user.branch, submitted=True  # Filter by BOTH branch AND submitted=True
        ).order_by('-submission_date')

    context = {
        'completed_audits': completed_audits,
        'incomplete_audits': incomplete_audits,
        'user_role': user.role,
    }
    return render(request, "audit_workflow/dashboard.html", context)


@login_required
def audit_detail(request, audit_id):
    audit = get_object_or_404(AuditSubmission, pk=audit_id)
    audit_objections = audit.auditobjection_set.all()

    context = {
        'audit': audit,
        'audit_objections': audit_objections,
    }
    return render(request, 'audit_workflow/audit_detail.html', context)


import logging
logger = logging.getLogger(__name__)



@login_required
def submit_objection(request, audit_id, item_id):
    audit = get_object_or_404(AuditSubmission, pk=audit_id)
    # Assuming you have a way to identify the specific item
    item = get_object_or_404(Items, id=item_id)

    existing_objection = AuditObjection.objects.filter(submission=audit, items=item).first()
    if existing_objection:
        logger.warning(f"Objection already exists for audit_id={audit_id}, item_id={item_id}")
        return JsonResponse({"error": "Objection for this item has already been submitted!"}, status=400)

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    if request.method == 'POST':
        logger.info(f"AJAX POST request received for audit_id={audit_id}, item_id={item_id}")
        logger.info("Processing POST data...")
        logger.info(f"POST data: {request.POST}")
        logger.info(f"FILES data: {request.FILES}")

        objection_form = AuditObjectionForm(request.POST)
        decision_form = Personel_DecisionForm(request.POST, instance=Personel_DecisionForm().Meta.model()) # Initialize with an instance for potential saving
        upload_formset = UploadFormSet(request.POST, request.FILES, prefix='upload_set')
        person_formset = PersonFormSet(request.POST, prefix='person_set')

        if all([objection_form.is_valid(), decision_form.is_valid(), upload_formset.is_valid(), person_formset.is_valid()]): # Removed upload_formset.is_valid()
            logger.info("All forms are valid. Proceeding to save data.")
            objection = objection_form.save(commit=False)
            objection.submission = audit
            objection.items = item
            objection.save()
            print("AuditObjection saved:", objection)
            AuditItemStatus.objects.update_or_create(
                audit=audit, item=item,
                defaults={'is_completed': True}
            )
            #Save Personel_Decision (or update if exists)
            decision, created = Personel_Decision.objects.update_or_create(
                audit_id=audit,
                items=item,
                defaults={"isEmployee_involve": decision_form.cleaned_data["isEmployee_involve"]},
            )

            print("Personel_Decision saved/updated:", decision, "Created:", created)

            # Save Persons
            if decision.isEmployee_involve:
                for person_form in person_formset:
                    if person_form.is_valid() and person_form.has_changed():
                        person = person_form.save(commit=False)
                        person.audit_id = audit
                        person.items = item
                        person.audit_objection = objection
                        person.save()
                        print("Person saved:", person)


            # Handle uploads (temporarily skipped)
            for upload_form in upload_formset:
                if upload_form.is_valid() and upload_form.cleaned_data:
                    upload = upload_form.save(commit=False)
                    upload.upload_link = objection
                    upload.save()
                    print("Upload saved:", upload)

            return JsonResponse({'success': 'Objection submitted successfully!'})
        else:
            logger.warning("Form validation failed.")
            errors = {}
            if objection_form.errors:
                errors['objection_form'] = objection_form.errors.as_json()
                logger.warning(f"Objection Form Errors: {objection_form.errors.as_json()}")
            if decision_form.errors:
                errors['decision_form'] = decision_form.errors.as_json()
                logger.warning(f"Decision Form Errors: {decision_form.errors.as_json()}")
            if upload_formset.errors:
                upload_errors = {}
                for i, errors_in_form in enumerate(upload_formset.errors):
                    if errors_in_form:
                        upload_errors[f'form-{i}'] = errors_in_form.as_json()
                errors['upload_formset'] = upload_errors
                logger.warning(f"Upload Formset Errors: {upload_errors}")
            if person_formset.errors:
                person_errors = {}
                for i, errors_in_form in enumerate(person_formset.errors):
                    if errors_in_form:
                        person_errors[f'form-{i}'] = errors_in_form.as_json()
                errors['person_formset'] = person_errors
                logger.warning(f"Person Formset Errors: {person_errors}")
            error_response = {"error": "Form validation failed", "errors": errors}
            logger.warning(f"Returning JSON error response: {error_response}")
            return JsonResponse(error_response, status=400)

    else:
        objection_form = AuditObjectionForm()
        decision_form = Personel_DecisionForm(instance=Personel_DecisionForm().Meta.model())
        upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set')
        person_formset = PersonFormSet(queryset=Person.objects.none(), prefix='person_set')

    return render(request, 'audit_workflow/submit_objection_form.html', {
        'objection_form': objection_form,
        'upload_formset': upload_formset, # Temporarily removed
        'decision_form': decision_form,
        'person_formset': person_formset,
        'audit': audit,
        'item': item,
        'upload_empty_form': upload_formset.empty_form    })


def edit_objection(request, objection_id):
    objection = get_object_or_404(AuditObjection, id=objection_id)
    audit = objection.submission
    item = objection.items

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    # Retrieve decision if exists (per audit+item)
    decision_instance = Personel_Decision.objects.filter(audit_id=audit, items=item).first()

    if request.method == "POST":
        objection_form = AuditObjectionForm(request.POST, instance=objection)
        decision_form = Personel_DecisionForm(request.POST, instance=decision_instance)

        person_formset = PersonFormSet(
            request.POST,
            queryset=Person.objects.filter(audit_id=audit, items=item)
        )
        upload_formset = UploadFormSet(
            request.POST, request.FILES,
            queryset=Upload.objects.filter(upload_link=objection),
            prefix="upload_set"
        )
        print("🔢 TOTAL_FORMS raw:", request.POST.get("person_set-TOTAL_FORMS"))
        print("🔢 TOTAL_FORMS (counted):", person_formset.total_form_count())
        if all([
            objection_form.is_valid(),
            decision_form.is_valid(),
            person_formset.is_valid(),
            upload_formset.is_valid()
        ]):
            # Save the objection
            objection = objection_form.save()

            # Save or update the decision
            decision, _ = Personel_Decision.objects.update_or_create(
                audit_id=audit,
                items=item,
                defaults={
                    "isEmployee_involve": decision_form.cleaned_data.get("isEmployee_involve")
                }
            )

            # Save related persons
            for person in person_formset.save(commit=False):
                person.audit_id = audit
                person.items = item
                person.audit_objection = objection
                person.save()
            for obj in person_formset.deleted_objects:
                obj.delete()

            # Save uploads
            for file in upload_formset.save(commit=False):
                if file.document:
                    file.upload_link = objection
                    file.save()
            for obj in upload_formset.deleted_objects:
                obj.delete()

            return redirect("audit_workflow:audit_detail", audit.id)

        else:
            print("❌ Form validation failed:", {
                "objection": objection_form.errors,
                "decision": decision_form.errors,
                "person_formset": person_formset.errors,
                "upload_formset": upload_formset.errors,
            })

    else:
        # GET method: prepare forms with prefilled data
        upload_queryset = Upload.objects.filter(upload_link=objection)
        person_queryset = Person.objects.filter(audit_id=audit, items=item)

        objection_form = AuditObjectionForm(instance=objection)
        decision_form = Personel_DecisionForm(instance=decision_instance)
        person_formset = PersonFormSet(queryset=person_queryset)
        upload_formset = UploadFormSet(queryset=upload_queryset, prefix="upload_set")

    return render(request, "audit_workflow/edit_objection.html", {
        "objection_form": objection_form,
        "decision_form": decision_form,
        "person_formset": person_formset,
        "upload_formset": upload_formset,
        "objection": objection,
        "audit": audit,
        "item": item,
    })


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, email=email, password=password)  # Authenticate using email
            if user is not None:
                login(request, user)
                return redirect("audit_workflow:start_audit")  # Redirect to your home page after login
            else:
                form.add_error(None, "Invalid email or password")  # Add error message to the form
    else:
        form = LoginForm()
    return render(request, 'audit_workflow/login.html', {'form': form})

def comment_detail(request, audit_id, item_id):
    audit = get_object_or_404(AuditSubmission, pk=audit_id)
    item = get_object_or_404(Items, pk=item_id)
    objection = AuditObjection.objects.filter(items=item, submission=audit).first()

    user_comment = Comment.objects.filter(commented_by=request.user, item=item, submission=audit, parent=None).first()  # Corrected query
    comments = Comment.objects.filter(item=item, submission=audit).order_by('-comment_date')  # Filter comments by submission too

    if request.method == "POST":
        form = CommentForm(request.POST, request.FILES)
        if form.is_valid():
            comment = form.save(commit=False)
            comment.commented_by = request.user
            comment.item = item
            comment.submission = audit

            if not user_comment:  # Now correctly checks for existing comment in this submission
                comment.save()
                return redirect('audit_workflow:comment_detail', audit_id=audit.id, item_id=item.id)

            parent_id = request.POST.get('parent_id')
            if parent_id:
                parent_comment = get_object_or_404(Comment, pk=parent_id)
                if not parent_comment.has_admin_manager_reply():
                    return JsonResponse({'success': False, 'error': "You can't reply until an admin/manager has responded."})

                comment.parent = parent_comment
                comment.save()
                return redirect('audit_workflow:comment_detail', audit_id=audit.id, item_id=item.id)

    else:
        form = CommentForm()

    context = {
        'audit': audit,
        'item': item,
        'objection': objection,
        'comments': comments,
        'form': form,
        'user_comment': user_comment,
    }

    return render(request, 'audit_workflow/comment_detail.html', context)


def user_logout(request):
    logout(request)
    return redirect('audit_workflow:login')

@login_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(user=request.user, data=request.POST)
        if form.is_valid():
            user = form.save()
            # Important: Update the session auth hash after password change
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password was successfully updated!')
            return redirect('audit_workflow:start_audit')  # Redirect to a success page or wherever you want
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PasswordChangeForm(user=request.user)
    return render(request, 'audit_workflow/change_password.html', {'form': form})

@login_required
def audit_report_page(request, audit_id):
    # Fetch the audit submission
    audit = get_object_or_404(AuditSubmission, id=audit_id)

    # Fetch related items for this audit submission
    items = Items.objects.filter(auditobjection__submission=audit).distinct()

    # Fetch objections for each item
    audit_data = []
    items = Items.objects.filter(auditobjection__submission=audit).distinct()

    # Fetch objections and comments for each item
    audit_data = []
    for item in items:
        objections = AuditObjection.objects.filter(submission=audit, items=item)
        comments = Comment.objects.filter(submission=audit, item=item).select_related('commented_by')

        audit_data.append({
            "item": item,
            "objections": objections,
            "comments": comments
        })

    context = {
        "audit": audit,
        "audit_data": audit_data,
    }

    return render(request, 'audit_workflow/audit_report_page.html', context)


import re
from django.utils.html import strip_tags  # Removes HTML tags
from html.parser import HTMLParser  # Converts entities like &nbsp; to spaces

class HTMLTextConverter(HTMLParser):
    """Convert HTML entities like &nbsp; to plain text."""
    def handle_data(self, data):
        self.text = data

    def get_text(self, html):
        self.text = ""
        self.feed(html)
        return self.text.strip()

def generate_audit_report(request, audit_id):
    # Fetch the audit submission
    audit = get_object_or_404(AuditSubmission, id=audit_id)
    doc = Document()
    doc.add_heading(f'Audit Report - {audit_id}', level=1)

    # General audit details
    doc.add_paragraph(f"Branch: {audit.branch.name}")
    doc.add_paragraph(f"Audit Period: {audit.start_date} to {audit.end_date}")
    doc.add_paragraph(f"Final Submission Date: {audit.submission_date}\n")

    items = Items.objects.filter(auditobjection__submission=audit).distinct()
    html_parser = HTMLTextConverter()  # Convert HTML entities

    for index, item in enumerate(items, start=1):
        doc.add_heading(f'{index}. {item.itemName}', level=2)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Amount(৳)"
        hdr_cells[2].text = "Comments"

        objections = AuditObjection.objects.filter(submission=audit, items=item)

        if objections.exists():
            for objection in objections:
                row_cells = table.add_row().cells
                row_cells[0].text = strip_tags(html_parser.get_text(objection.description))  # Strip HTML
                row_cells[1].text = f"{objection.amount}"

                # Fetch related comments
                comments = Comment.objects.filter(submission=audit, item=item)
                if comments.exists():
                    comment_text = "\n".join(
                        [f"{c.commented_by}: {strip_tags(html_parser.get_text(c.comment))}" for c in comments]
                    )
                else:
                    comment_text = "No comments"

                row_cells[2].text = comment_text
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "No objections found"
            row_cells[1].text = "-"
            row_cells[2].text = "-"

        doc.add_paragraph("")  # Spacing between tables

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="Audit_Report_{audit_id}.docx"'
    doc.save(response)

    return response


def add_or_update_item(request, item_id=None):
    # If item_id is provided, we are updating an existing item
    if item_id:
        item = get_object_or_404(Items, id=item_id)
    else:
        item = None

    if request.method == 'POST':
        form = ItemForm(request.POST, instance=item)
        if form.is_valid():
            form.save()
            messages.success(request, "Item saved successfully!")
            return redirect('audit_workflow:item_list')  # Redirect to item list or another page
        else:
            messages.error(request, "There was an error in the form.")
    else:
        form = ItemForm(instance=item)

    return render(request, 'audit_workflow/add_or_update_item.html', {'form': form, 'item': item})

def item_list(request):
    items = Items.objects.all()
    return render(request, 'audit_workflow/item_list.html', {'items': items})

def delete_item(request, item_id):
    item = get_object_or_404(Items, id=item_id)
    item.delete()
    messages.success(request, "Item deleted successfully!")
    return redirect('audit_workflow:item_list')
