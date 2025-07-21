from django.contrib.auth import authenticate, login
from django.core.paginator import Paginator
from django.forms import modelformset_factory, inlineformset_factory, formset_factory
from django.utils.timezone import now
from .forms import AuditSubmissionForm, CommentForm, Personel_DecisionForm, PersonForm, ItemForm, ObjectionDecisionForm
from .models import AuditSubmission, Items, AuditObjection, AuditItemStatus, Branch, User, Person, Comment, Upload, \
    Personel_Decision, ObjectionDecision
from .forms import AuditObjectionForm, LoginForm, UploadForm
from django.contrib.auth.decorators import login_required, user_passes_test
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
from django.contrib import messages
from django.db import IntegrityError
from django.contrib.auth import logout, update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm
import re
from django.utils.html import strip_tags  # Removes HTML tags
from html.parser import HTMLParser  # Converts entities like &nbsp; to space
from docx import Document

from django.http import HttpResponseForbidden
from .models import AuditTypeChoices
#from ..bangaidj.user_creation import branch


@login_required
def start_audit(request):
    if not request.user.is_staff and not request.user.is_superuser:
        messages.error(request, "You are not authorized to start a new audit.")
        return redirect('audit_workflow:dashboard')

    if request.method == "POST":
        form = AuditSubmissionForm(request.POST)
        all_branches = Branch.objects.all().distinct()

        if form.is_valid():
            audit = form.save(commit=False)
            audit.auditor = request.user
            audit.save()

            # Select only Items matching the selected audit_type
            selected_audit_type = audit.audit_type  # Get the audit_type from the form
            items = Items.objects.filter(audit_type=selected_audit_type)

            for item in items:
                AuditItemStatus.objects.create(audit=audit, item=item)

            return redirect('audit_workflow:audit_summary', audit_id=audit.id)

    else:
        form = AuditSubmissionForm()
        all_branches = Branch.objects.all().distinct()

    return render(request, 'audit_workflow/start_audit.html', {
        'form': form,
        'all_branches': all_branches,
    })


@login_required()
def audit_summary(request, audit_id):
    try:
        audit = AuditSubmission.objects.get(id=audit_id, auditor=request.user)
    except AuditSubmission.DoesNotExist:
        return render(request, 'audit_workflow/404.html', status=404)

    items = Items.objects.filter(audit_type=audit.audit_type).order_by('id')
    paginator = Paginator(items, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    item_statuses = {status.item.id: status.is_completed for status in AuditItemStatus.objects.filter(audit=audit)}
    selected_item = None
    form_html = ""

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    convicted_persons_by_item = {
        item.id: Person.objects.filter(audit_id=audit, items=item)
        for item in items
    }

    if request.method == "GET" and 'item_id' in request.GET:
        selected_item_id = request.GET.get('item_id')
        selected_item = get_object_or_404(Items, id=selected_item_id)
        objection_form = AuditObjectionForm(initial={'items': selected_item})
        upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set')  # Set prefix here
        decision_form = Personel_DecisionForm()
        person_formset = PersonFormSet(queryset=Person.objects.filter(audit_id=audit, items=selected_item), prefix='person_set') # Set prefix here

        form_html = render_to_string(
            "audit_workflow/submit_objection_form.html",
            {
                "objection_form": objection_form,
                "upload_formset": upload_formset,  # Pass UploadFormSet to the template
                "audit": audit,
                "item": selected_item,
                "decision_form": decision_form,
                "person_formset": person_formset,
            },
            request=request,
        )
        return HttpResponse(form_html)

    # Initialize forms for the main context
    decision_form = Personel_DecisionForm()
    person_formset = PersonFormSet(queryset=Person.objects.none(), prefix='person_set') # Set prefix here
    upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set') # Set prefix here

    context = {
        "audit": audit,
        "items": page_obj,
        "item_statuses": item_statuses,
        "selected_item": selected_item,
        "form_html": form_html,
        "decision_form": decision_form,
        "person_formset": person_formset,
        'page_obj': page_obj,
        "upload_formset": upload_formset, # Pass UploadFormSet to the main context
        "convicted_persons_by_item": convicted_persons_by_item,
    }
    return render(request, "audit_workflow/audit_summary.html", context)






@login_required
def finalize_audit(request, audit_id):
    audit = get_object_or_404(AuditSubmission, id=audit_id, auditor=request.user)
    incomplete_items = AuditItemStatus.objects.filter(audit=audit, is_completed=False)

    if incomplete_items.exists():
        return render(request, "audit_workflow/finalize_error.html", {"audit": audit})

    audit.submitted = True
    audit.submission_date = now()
    audit.save()
    return redirect("audit_workflow:dashboard")



def dashboard(request):
    user = request.user
    completed_audits = []
    incomplete_audits = []

    if user.role == 'admin':
        all_audits = AuditSubmission.objects.all().order_by('-submission_date')

        for audit in all_audits:
            all_items_completed = all(status.is_completed for status in AuditItemStatus.objects.filter(audit=audit))
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)


    elif user.role == 'authorizer':
        all_region_audits = AuditSubmission.objects.none()

        # ✅ Regular authorizer: only audits in their division
        division_audits = AuditSubmission.objects.filter(
            branch__division__iexact=user.division
        )

        # ✅ Special authorizer (DGM Compliance): audits done by RAO and auditor is from 'Head Office'
        head_office_rao_audits = AuditSubmission.objects.filter(
            auditor__role='rao',
            auditor__division='Head Office'
        )

        # Combine both querysets (removes duplicates if any)
        all_region_audits = (division_audits | head_office_rao_audits).distinct().order_by('-submission_date')

        for audit in all_region_audits:
            all_items_completed = all(
                status.is_completed for status in AuditItemStatus.objects.filter(audit=audit)
            )
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)

    elif user.role == 'manager':
        # Filter AuditSubmissions where the region of the associated branch matches the manager's region
        all_region_audits = AuditSubmission.objects.filter(branch__division__iexact=user.division).order_by('-submission_date')

        for audit in all_region_audits:
            all_items_completed = all(
                status.is_completed for status in AuditItemStatus.objects.filter(audit=audit)
            )
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)
    elif user.role == 'rao' and user.division == 'Head Office':
        all_audits = AuditSubmission.objects.all().order_by('-submission_date')

        for audit in all_audits:
            all_items_completed = all(status.is_completed for status in AuditItemStatus.objects.filter(audit=audit))
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)

    elif user.role == 'rao':
        # Filter AuditSubmissions where the region of the associated branch matches the manager's region
        all_region_audits = AuditSubmission.objects.filter(branch__region__iexact=user.region).order_by('-submission_date')

        for audit in all_region_audits:
            all_items_completed = all(
                status.is_completed for status in AuditItemStatus.objects.filter(audit=audit)
            )
            if all_items_completed:
                completed_audits.append(audit)
            else:
                incomplete_audits.append(audit)

    else:  # Regular user
        completed_audits = AuditSubmission.objects.filter(
            branch=user.branch, submitted=True
        ).order_by('-submission_date')

    context = {
        'completed_audits': completed_audits,
        'incomplete_audits': incomplete_audits,
        'user_role': user.role,
    }
    return render(request, "audit_workflow/dashboard.html", context)


@login_required
def audit_detail(request, audit_id):
    audit = get_object_or_404(AuditSubmission, pk=audit_id)
    audit_objections = audit.auditobjection_set.select_related('decision', 'items').all()

    context = {
        'audit': audit,
        'audit_objections': audit_objections,
            }
    return render(request, 'audit_workflow/audit_detail.html', context)


import logging
logger = logging.getLogger(__name__)



@login_required
def submit_objection(request, audit_id, item_id):
    audit = get_object_or_404(AuditSubmission, pk=audit_id)
    # Assuming you have a way to identify the specific item
    item = get_object_or_404(Items, id=item_id)

    existing_objection = AuditObjection.objects.filter(submission=audit, items=item).first()
    if existing_objection:
        logger.warning(f"Objection already exists for audit_id={audit_id}, item_id={item_id}")
        return JsonResponse({"error": "Objection for this item has already been submitted!"}, status=400)

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    if request.method == 'POST':
        logger.info(f"AJAX POST request received for audit_id={audit_id}, item_id={item_id}")
        logger.info("Processing POST data...")
        logger.info(f"POST data: {request.POST}")
        logger.info(f"FILES data: {request.FILES}")

        objection_form = AuditObjectionForm(request.POST)
        decision_form = Personel_DecisionForm(request.POST, instance=Personel_DecisionForm().Meta.model()) # Initialize with an instance for potential saving
        upload_formset = UploadFormSet(request.POST, request.FILES, prefix='upload_set')
        person_formset = PersonFormSet(request.POST, prefix='person_set')

        if all([objection_form.is_valid(), decision_form.is_valid(), upload_formset.is_valid(), person_formset.is_valid()]): # Removed upload_formset.is_valid()
            logger.info("All forms are valid. Proceeding to save data.")
            objection = objection_form.save(commit=False)
            objection.submission = audit
            objection.items = item
            objection.save()
            print("AuditObjection saved:", objection)
            AuditItemStatus.objects.update_or_create(
                audit=audit, item=item,
                defaults={'is_completed': True}
            )
            #Save Personel_Decision (or update if exists)
            decision, created = Personel_Decision.objects.update_or_create(
                audit_id=audit,
                items=item,
                defaults={"isEmployee_involve": decision_form.cleaned_data["isEmployee_involve"]},
            )

            print("Personel_Decision saved/updated:", decision, "Created:", created)

            # Save Persons
            if decision.isEmployee_involve:
                for person_form in person_formset:
                    if person_form.is_valid() and person_form.has_changed():
                        person = person_form.save(commit=False)
                        person.audit_id = audit
                        person.items = item
                        person.audit_objection = objection
                        person.save()
                        print("Person saved:", person)


            # Handle uploads (temporarily skipped)
            for upload_form in upload_formset:
                if upload_form.is_valid() and upload_form.cleaned_data:
                    upload = upload_form.save(commit=False)
                    upload.upload_link = objection
                    upload.save()
                    print("Upload saved:", upload)

            return JsonResponse({'success': 'Objection submitted successfully!'})
        else:
            logger.warning("Form validation failed.")
            errors = {}
            if objection_form.errors:
                errors['objection_form'] = objection_form.errors.as_json()
                logger.warning(f"Objection Form Errors: {objection_form.errors.as_json()}")
            if decision_form.errors:
                errors['decision_form'] = decision_form.errors.as_json()
                logger.warning(f"Decision Form Errors: {decision_form.errors.as_json()}")
            if upload_formset.errors:
                upload_errors = {}
                for i, errors_in_form in enumerate(upload_formset.errors):
                    if errors_in_form:
                        upload_errors[f'form-{i}'] = errors_in_form.as_json()
                errors['upload_formset'] = upload_errors
                logger.warning(f"Upload Formset Errors: {upload_errors}")
            if person_formset.errors:
                person_errors = {}
                for i, errors_in_form in enumerate(person_formset.errors):
                    if errors_in_form:
                        person_errors[f'form-{i}'] = errors_in_form.as_json()
                errors['person_formset'] = person_errors
                logger.warning(f"Person Formset Errors: {person_errors}")
            error_response = {"error": "Form validation failed", "errors": errors}
            logger.warning(f"Returning JSON error response: {error_response}")
            return JsonResponse(error_response, status=400)

    else:
        objection_form = AuditObjectionForm()
        decision_form = Personel_DecisionForm(instance=Personel_DecisionForm().Meta.model())
        upload_formset = UploadFormSet(queryset=Upload.objects.none(), prefix='upload_set')
        person_formset = PersonFormSet(queryset=Person.objects.none(), prefix='person_set')

    return render(request, 'audit_workflow/submit_objection_form.html', {
        'objection_form': objection_form,
        'upload_formset': upload_formset, # Temporarily removed
        'decision_form': decision_form,
        'person_formset': person_formset,
        'audit': audit,
        'item': item,
        'upload_empty_form': upload_formset.empty_form    })


def edit_objection(request, objection_id):
    objection = get_object_or_404(AuditObjection, id=objection_id)
    audit = objection.submission
    item = objection.items

    UploadFormSet = modelformset_factory(Upload, form=UploadForm, extra=1, can_delete=True)
    PersonFormSet = modelformset_factory(Person, form=PersonForm, extra=1, can_delete=True)

    # Retrieve decision if exists (per audit+item)
    decision_instance = Personel_Decision.objects.filter(audit_id=audit, items=item).first()

    if request.method == "POST":
        objection_form = AuditObjectionForm(request.POST, instance=objection)
        decision_form = Personel_DecisionForm(request.POST, instance=decision_instance)

        person_formset = PersonFormSet(
            request.POST,
            queryset=Person.objects.filter(audit_id=audit, items=item)
        )
        upload_formset = UploadFormSet(
            request.POST, request.FILES,
            queryset=Upload.objects.filter(upload_link=objection),
            prefix="upload_set"
        )
        print("🔢 TOTAL_FORMS raw:", request.POST.get("person_set-TOTAL_FORMS"))
        print("🔢 TOTAL_FORMS (counted):", person_formset.total_form_count())
        if all([
            objection_form.is_valid(),
            decision_form.is_valid(),
            person_formset.is_valid(),
            upload_formset.is_valid()
        ]):
            # Save the objection
            objection = objection_form.save()

            # Save or update the decision
            decision, _ = Personel_Decision.objects.update_or_create(
                audit_id=audit,
                items=item,
                defaults={
                    "isEmployee_involve": decision_form.cleaned_data.get("isEmployee_involve")
                }
            )

            # Save related persons
            for person in person_formset.save(commit=False):
                person.audit_id = audit
                person.items = item
                person.audit_objection = objection
                person.save()
            for obj in person_formset.deleted_objects:
                obj.delete()

            # Save uploads
            for file in upload_formset.save(commit=False):
                if file.document:
                    file.upload_link = objection
                    file.save()
            for obj in upload_formset.deleted_objects:
                obj.delete()

            return redirect("audit_workflow:audit_detail", audit.id)

        else:
            print("❌ Form validation failed:", {
                "objection": objection_form.errors,
                "decision": decision_form.errors,
                "person_formset": person_formset.errors,
                "upload_formset": upload_formset.errors,
            })

    else:
        # GET method: prepare forms with prefilled data
        upload_queryset = Upload.objects.filter(upload_link=objection)
        person_queryset = Person.objects.filter(audit_id=audit, items=item)

        objection_form = AuditObjectionForm(instance=objection)
        decision_form = Personel_DecisionForm(instance=decision_instance)
        person_formset = PersonFormSet(queryset=person_queryset)
        upload_formset = UploadFormSet(queryset=upload_queryset, prefix="upload_set")

    return render(request, "audit_workflow/edit_objection.html", {
        "objection_form": objection_form,
        "decision_form": decision_form,
        "person_formset": person_formset,
        "upload_formset": upload_formset,
        "objection": objection,
        "audit": audit,
        "item": item,
    })


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, email=email, password=password)  # Authenticate using email
            if user is not None:
                login(request, user)
                return redirect("audit_workflow:dashboard")  # Redirect to your home page after login
            else:
                form.add_error(None, "Invalid email or password")  # Add error message to the form
    else:
        form = LoginForm()
    return render(request, 'audit_workflow/login.html', {'form': form})


def user_logout(request):
    logout(request)
    return redirect('audit_workflow:login')

@login_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(user=request.user, data=request.POST)
        if form.is_valid():
            user = form.save()
            # Important: Update the session auth hash after password change
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password was successfully updated!')
            return redirect('audit_workflow:start_audit')  # Redirect to a success page or wherever you want
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PasswordChangeForm(user=request.user)
    return render(request, 'audit_workflow/change_password.html', {'form': form})

@login_required
def audit_report_page(request, audit_id):
    # Fetch the audit submission
    audit = get_object_or_404(AuditSubmission, id=audit_id)

    # Fetch related items for this audit submission
    items = Items.objects.filter(auditobjection__submission=audit).distinct()

    # Fetch objections for each item
    audit_data = []
    items = Items.objects.filter(auditobjection__submission=audit).distinct()

    # Fetch objections and comments for each item
    audit_data = []
    for item in items:
        objections = AuditObjection.objects.filter(submission=audit, items=item)
        comments = Comment.objects.filter(submission=audit, item=item).select_related('commented_by')

        audit_data.append({
            "item": item,
            "objections": objections,
            "comments": comments
        })

    context = {
        "audit": audit,
        "audit_data": audit_data,
    }

    return render(request, 'audit_workflow/audit_report_page.html', context)


import re
from django.utils.html import strip_tags  # Removes HTML tags
from html.parser import HTMLParser  # Converts entities like &nbsp; to spaces

class HTMLTextConverter(HTMLParser):
    """Convert HTML entities like &nbsp; to plain text."""
    def handle_data(self, data):
        self.text = data

    def get_text(self, html):
        self.text = ""
        self.feed(html)
        return self.text.strip()

def generate_audit_report(request, audit_id):
    # Fetch the audit submission
    audit = get_object_or_404(AuditSubmission, id=audit_id)
    doc = Document()
    doc.add_heading(f'Audit Report - {audit_id}', level=1)

    # General audit details
    doc.add_paragraph(f"Branch: {audit.branch.name}")
    doc.add_paragraph(f"Audit Period: {audit.start_date} to {audit.end_date}")
    doc.add_paragraph(f"Final Submission Date: {audit.submission_date}\n")

    items = Items.objects.filter(auditobjection__submission=audit).distinct()
    html_parser = HTMLTextConverter()  # Convert HTML entities

    for index, item in enumerate(items, start=1):
        doc.add_heading(f'{index}. {item.itemName}', level=2)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Amount(৳)"
        hdr_cells[2].text = "Comments"

        objections = AuditObjection.objects.filter(submission=audit, items=item)

        if objections.exists():
            for objection in objections:
                row_cells = table.add_row().cells
                row_cells[0].text = strip_tags(html_parser.get_text(objection.description))  # Strip HTML
                row_cells[1].text = f"{objection.amount}"

                # Fetch related comments
                comments = Comment.objects.filter(submission=audit, item=item)
                if comments.exists():
                    comment_text = "\n".join(
                        [f"{c.commented_by}: {strip_tags(html_parser.get_text(c.comment))}" for c in comments]
                    )
                else:
                    comment_text = "No comments"

                row_cells[2].text = comment_text
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "No objections found"
            row_cells[1].text = "-"
            row_cells[2].text = "-"

        doc.add_paragraph("")  # Spacing between tables

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="Audit_Report_{audit_id}.docx"'
    doc.save(response)

    return response


def add_or_update_item(request, item_id=None):
    # If item_id is provided, we are updating an existing item
    if item_id:
        item = get_object_or_404(Items, id=item_id)
    else:
        item = None

    if request.method == 'POST':
        form = ItemForm(request.POST, instance=item)
        if form.is_valid():
            form.save()
            messages.success(request, "Item saved successfully!")
            return redirect('audit_workflow:item_list')  # Redirect to item list or another page
        else:
            messages.error(request, "There was an error in the form.")
    else:
        form = ItemForm(instance=item)

    return render(request, 'audit_workflow/add_or_update_item.html', {'form': form, 'item': item})

def item_list(request):
    items = Items.objects.all()
    return render(request, 'audit_workflow/item_list.html', {'items': items})

def delete_item(request, item_id):
    item = get_object_or_404(Items, id=item_id)
    item.delete()
    messages.success(request, "Item deleted successfully!")
    return redirect('audit_workflow:item_list')

def is_admin(user):
    return user.role == 'admin'


@login_required
@user_passes_test(lambda u: u.is_superuser or u.is_staff or u.role=='manager')
def admin_objection_decision(request, audit_objection_id):
    audit_objection = get_object_or_404(AuditObjection, id=audit_objection_id)
    decision, created = ObjectionDecision.objects.get_or_create(audit_objection=audit_objection)

    if request.method == 'POST':
        form = ObjectionDecisionForm(request.POST, instance=decision)
        if form.is_valid():
            new_status = form.cleaned_data.get('status')

            # Only superusers can reopen (status = 'run')
            if audit_objection.decision == 'closed' and new_status != 'closed' and not request.user.is_superuser:
                messages.error(request, "Only superusers can reopen a closed objection.")
                return redirect('audit_workflow:admin_objection_decision', audit_objection_id=audit_objection.id)

            decision = form.save(commit=False)
            decision.decided_by = request.user
            decision.save()

            audit_objection.status = new_status
            audit_objection.save()

            messages.success(request, "Objection status updated.")
            return redirect('audit_workflow:audit_detail', audit_id=audit_objection.submission.id)
    else:
        form = ObjectionDecisionForm(instance=decision)

    return render(request, 'audit_workflow/admin_objection_decision.html', {
        'form': form,
        'audit_objection': audit_objection,
    })

def get_effective_user_role(user, submission):
    # If the audit is on a manager's own branch
    if user.role == 'manager' and submission.branch == user.branch:
        return 'regular'  # Manager is now acting as audited person
    if user.role == 'rao' and submission.branch == user.branch:
        return 'regular'  # rao is now acting as audited person
    # If the audit is started by an admin, and the admin is commenting
    if user.role == 'admin' and submission.auditor.role == 'rao':
        return 'manager'  # Admin is acting as the auditing manager

    return user.role




@login_required
def comment_detail(request, audit_id, item_id):
    item = get_object_or_404(Items, id=item_id)
    submission = get_object_or_404(AuditSubmission, id=audit_id)
    objection = AuditObjection.objects.filter(submission=submission, items=item).first()
    comments = Comment.objects.filter(submission=submission, item=item).order_by('comment_date')
    user = request.user
    user_role = get_effective_user_role(user, submission)
    is_closed = objection.decision.decision == 'close' if objection and hasattr(objection, 'decision') else False

    all_items = list(Items.objects.all().order_by("id"))
    item_ids = [i.id for i in all_items]
    index = item_ids.index(item.id)
    next_item_id = item_ids[index + 1] if index + 1 < len(item_ids) else None
    previous_item_id = item_ids[index - 1] if index - 1 >= 0 else None

    comment_form = CommentForm()
    reply_forms = {}
    can_reply_dict = {}
    can_comment = False
    can_edit_comment = False
    editable_comment = None

    # Find top-level and latest comments by user
    user_comments = comments.filter(commented_by=user).order_by('-comment_date')
    for c in user_comments:
        if not c.replies.exists():
            editable_comment = c
            can_edit_comment = True
            break

    # First-time comment check
    if user_role == 'regular' and not is_closed and not comments.filter(commented_by=user, parent=None).exists():
        can_comment = True

    top_level_comment = comments.filter(parent=None).first()
    predefined_reply = [
        "নিস্পত্তি যোগ্য নয়।",
        "আদায় করে অবহিত করুন",
        "সম্বন্নয় / নিয়মিত করে অবহিত করুন",
        "পরবর্তী নিরীক্ষা যাচাই স্বাপেক্ষে নিস্পত্তি করা হবে।",
        "শাখার জবাবের প্রেক্ষিতে নিস্পত্তি।"
    ]

    if top_level_comment and not is_closed:
        latest_comment = comments.last()
        if latest_comment.commented_by != user:
            if (
                get_effective_user_role(latest_comment.commented_by, submission) == 'manager'
                and user_role == 'regular'
            ) or (
                get_effective_user_role(latest_comment.commented_by, submission) == 'regular'
                and user_role == 'manager'
            ):
                can_reply = True

    for comment in comments.filter(parent=None):
        latest = comment.replies.last() if comment.replies.exists() else comment
        can_reply_dict[comment.id] = False
        if not is_closed and latest.commented_by != user:
            latest_role = get_effective_user_role(latest.commented_by, submission)
            if (
                latest_role == 'manager' and user_role == 'regular'
            ) or (
                latest_role == 'regular' and user_role == 'manager'
            ):
                can_reply_dict[comment.id] = True
        prefix = f"comment_{comment.id}"
        reply_forms[comment.id] = CommentForm(prefix=prefix, initial={'parent_id': comment.id})

        if request.method == "POST" and "edit_comment" in request.POST:
            comment_id = request.POST.get("edit_comment")
            new_text = request.POST.get("edit_text")
            comment = get_object_or_404(Comment, id=comment_id, commented_by=request.user)
            if not comment.replies.exists():
                comment.comment = new_text
                comment.save()
                return redirect("audit_workflow:comment_detail", audit_id=submission.id, item_id=item.id)
        elif 'comment_form' in request.POST and can_comment:
            comment_form = CommentForm(request.POST, request.FILES)
            if comment_form.is_valid():
                new_comment = comment_form.save(commit=False)
                new_comment.item = item
                new_comment.submission = submission
                new_comment.commented_by = user
                new_comment.save()
                submission.is_authorized = False
                submission.save()
                if 'save_and_next' in request.POST and next_item_id:
                    return redirect('audit_workflow:comment_detail', audit_id=audit_id, item_id=next_item_id)
                return redirect('audit_workflow:comment_detail', audit_id=audit_id, item_id=item.id)

        elif 'reply_submit' in request.POST:
            comment_id = request.POST.get('reply_submit')
            if comment_id and comment_id.isdigit() and int(comment_id) in can_reply_dict and can_reply_dict[int(comment_id)]:
                prefix = f"comment_{comment_id}"
                form = CommentForm(request.POST, request.FILES, prefix=prefix)
                if form.is_valid():
                    parent_id = form.cleaned_data['parent_id']
                    parent_comment = get_object_or_404(Comment, id=parent_id)
                    reply = form.save(commit=False)
                    reply.parent = parent_comment
                    reply.item = item
                    reply.submission = submission
                    reply.commented_by = user
                    reply.comment = (
                        request.POST.get(f'{prefix}-predefined_reply') or form.cleaned_data.get('comment', '')
                    )
                    reply.save()
                    submission.is_authorized = False
                    submission.save()
                    return redirect('audit_workflow:comment_detail', audit_id=audit_id, item_id=item.id)

    context = {
        'item': item,
        'submission': submission,
        'objection': objection,
        'comments': comments,
        'comment_form': comment_form,
        'reply_forms': reply_forms,
        'can_reply_dict': can_reply_dict,
        'can_comment': can_comment,
        'can_edit_comment': can_edit_comment,
        'editable_comment': editable_comment,
        'top_level_comment': top_level_comment,
        'is_closed': is_closed,
        'is_manager': user_role == 'manager',
        'is_regular': user_role == 'regular',
        'predefined_reply': predefined_reply,
        'next_item_id': next_item_id,
        'previous_item_id': previous_item_id,
    }
    return render(request, 'audit_workflow/comment_detail.html', context)




def get_effective_manager_comments(submission):
    """
    Fetch latest manager replies for each item in the submission.
    """
    comments = (
        Comment.objects
        .filter(submission=submission)
        .select_related('commented_by', 'item', 'parent')
        .order_by('item_id', '-comment_date')
    )

    latest_by_item = {}
    for comment in comments:
        if comment.item_id not in latest_by_item:
            role = get_effective_user_role(comment.commented_by, submission)
            if role == 'manager':
                latest_by_item[comment.item_id] = comment

    return list(latest_by_item.values())



from django.shortcuts import render, get_object_or_404
from collections import defaultdict
from .models import AuditSubmission, Comment, Items
import html

@login_required
def jaripotro_report(request, audit_id):
    submission = get_object_or_404(AuditSubmission, id=audit_id)

    # Step 1: Fetch all top-level and reply comments made by manager for this submission
    # manager_comments = (
    #     Comment.objects
    #     .filter(submission=submission, commented_by__role='manager')
    #     .select_related('item')  # Optimizes access to comment.item
    #     .order_by('item_id', '-comment_date')
    # )
    manager_comments = get_effective_manager_comments(submission)
    # Step 2: Track the latest manager comment per item
    latest_comments_by_item = {}
    for comment in manager_comments:
        item_id = comment.item_id
        if item_id not in latest_comments_by_item:
            latest_comments_by_item[item_id] = comment

    # Step 3: Group comments by comment text and collect itemNos
    grouped = defaultdict(list)
    for item_id, comment in latest_comments_by_item.items():
        item_no = comment.item.itemNo  # assumes `itemNo` exists in `Items`
        #full_text = comment.comment.strip() if comment.comment else "অনির্ধারিত মন্তব্য"
        full_text = html.unescape(strip_tags(comment.comment.strip())) if comment.comment else "অনির্ধারিত মন্তব্য"
        grouped[full_text].append(item_no)

    # Step 4: Build rows
    report_data = []
    for comment_text, item_nos in grouped.items():
        item_nos.sort()
        report_data.append({
            'item_nos': ", ".join(str(no) for no in item_nos),
            'comment_text': comment_text
        })

    context = {
        'submission': submission,
        'report_data': report_data
    }
    return render(request, 'audit_workflow/jaripotro_report.html', context)


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict
from django.utils.html import strip_tags
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
import os
import html
from docx.shared import RGBColor
from docx.oxml.ns import qn
bangla_digits = {
    '0': '০', '1': '১', '2': '২', '3': '৩', '4': '৪',
    '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯'
}

def to_bangla_date(date_obj):
    date_str = date_obj.strftime('%d-%m-%Y')
    return ''.join(bangla_digits.get(ch, ch) for ch in date_str)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_bottom_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # Thickness
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)
    tcPr.append(borders)

@login_required
@user_passes_test(lambda u: u.is_superuser or u.is_staff or u.role == 'manager' or u.role == 'regular')
def download_jaripotro_docx(request, audit_id):
    submission = get_object_or_404(AuditSubmission, id=audit_id)

    if not submission.is_authorized and request.user.role != 'authorizer':
        return HttpResponseForbidden("You are not allowed to download this report until it is authorized.")

    # Step 1: Get all manager comments
    # manager_comments = (
    #     Comment.objects
    #     .filter(submission=submission, commented_by__role='manager')
    #     .select_related('item')
    #     .order_by('item_id', '-comment_date')
    # )
    manager_comments = get_effective_manager_comments(submission)

    # Step 2: Get latest manager comment per item
    latest_comments_by_item = {}
    for comment in manager_comments:
        if comment.item_id not in latest_comments_by_item:
            latest_comments_by_item[comment.item_id] = comment

    # Step 3: Group by comment text
    grouped = defaultdict(list)
    for item_id, comment in latest_comments_by_item.items():
        item_no = comment.item.itemNo if comment.item else "N/A"
        full_text = html.unescape(strip_tags(comment.comment.strip())) if comment.comment else "অনির্ধারিত মন্তব্য"
        grouped[full_text].append(item_no)

    # Step 4: Create DOCX and reduce top margin
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.4)   # Reduce top margin to 0.2 inches
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Nikosh'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    # Header table
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = True
    hdr_cells = header_table.rows[0].cells

    # Left: Logo
    logo_path = os.path.join('static', 'images', 'krishibank_logo.png')
    if os.path.exists(logo_path):
        paragraph = hdr_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(0.6))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Center: Bank info
    center_para = hdr_cells[1].paragraphs[0]
    run = center_para.add_run("বাংলাদেশ কৃষি ব্যাংক\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(3, 82, 3)
    user = request.user
    center_para.add_run({submission.authorized_by.branch.bn_name})
    # center_para.add_run(f"অঞ্চলঃ{submission.authorized_by.branch.region}\n")
    # center_para.add_run(f"বিভাগঃ{submission.authorized_by.branch.division}\n")
    center_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Right: Contact
    right_para = hdr_cells[2].paragraphs[0]
    right_para.add_run(f"ইমেইলঃ{submission.authorized_by.email}\n")
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Apply bottom border to all cells in the row
    for cell in hdr_cells:
        set_bottom_border(cell)
    # under line:
    para = doc.add_paragraph()


    # Manager identity
    manager_para = doc.add_paragraph()
    manager_para.paragraph_format.space_before = Pt(0)
    manager_para.paragraph_format.space_after = Pt(6)
    manager_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run1 = manager_para.add_run("ব্যবস্থাপক\n")
    run2 = manager_para.add_run(f"{submission.branch.bn_name} \n")
    #run3 = manager_para.add_run(f"{submission.branch.division}\n")
    run4 = manager_para.add_run("বাংলাদেশ কৃষি ব্যাংক\n")

    for run in [run1, run2, run4]:
        run.font.name = 'Nikosh'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    # Heading
    start_bangla = to_bangla_date(submission.start_date)
    end_bangla = to_bangla_date(submission.end_date)
    heading_text = (
        f"বিষয়ঃ বাংলাদেশ কৃষি ব্যাংক, {submission.branch.bn_name} এর "
        f"{start_bangla} হতে {end_bangla} পর্যন্ত সময়ের "
        f"আভ্যন্তরীণ নিরীক্ষা প্রতিবেদনের অনুচ্ছেদওয়ারী পরিপালন জবাব প্রসংগে।"
    )

    heading_para = doc.add_paragraph()
    heading_para.paragraph_format.space_before = Pt(0)
    heading_para.paragraph_format.space_after = Pt(12)
    heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading_para.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Nikosh'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    # Comments Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True

    hdr_cells = table.rows[0].cells
    hdr_texts = ['অনুচ্ছেদ নং', 'অনুচ্ছেদওয়ারী মন্তব্য/প্রস্তাবনা']

    for i, text in enumerate(hdr_texts):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Nikosh'
        run.font.size = Pt(14)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    for comment_text, item_nos in grouped.items():
        row_cells = table.add_row().cells
        cell_texts = [", ".join(str(no) for no in sorted(item_nos)), comment_text]
        for i, text in enumerate(cell_texts):
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run(text)
            run.font.name = 'Nikosh'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    # Signature
    #     table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Move the table to the right side

    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table.autofit = True
    table.allow_autofit = True
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.5)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run1 = paragraph.add_run("উপমহাব্যবস্থাপক\n")
    run1.font.size = Pt(14)

    if submission.authorized_by:
        run2 = paragraph.add_run(f"{submission.authorized_by.branch.bn_name}")
    else:
        run2 = paragraph.add_run("অননুমোদিত")

    run2.font.size = Pt(12)

    for run in [run1, run2]:
        run.font.name = 'Nikosh'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nikosh')

    # Export
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = f'attachment; filename="Jaripotro_{audit_id}.docx"'
    doc.save(response)
    return response


# Supporting function
def get_predefined_text(obj):
    comment = obj.comments.filter(user__role='manager', predefined_text__isnull=False).order_by('-created_at').first()
    return comment.predefined_text if comment else ''

@login_required
@user_passes_test(lambda u: u.role == 'authorizer')
def authorize_submission(request, audit_id):
    submission = get_object_or_404(AuditSubmission, id=audit_id)
    submission.is_authorized = True
    submission.authorized_by = request.user
    submission.save()
    messages.success(request, "Submission has been authorized.")
    return redirect('audit_workflow:jaripotro_report', audit_id=audit_id)

from django.db.models.signals import post_save
from django.dispatch import receiver

@receiver(post_save, sender=Comment)
def revoke_authorization_on_manager_comment(sender, instance, **kwargs):
    if instance.commented_by.role == 'manager':
        submission = instance.submission
        submission.is_authorized = False
        submission.authorized_by = None
        submission.save()


